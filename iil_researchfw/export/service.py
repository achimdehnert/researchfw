"""Research export — Markdown, LaTeX, BibTeX, DOCX."""
from __future__ import annotations

import logging
from datetime import datetime
from typing import Any

from iil_researchfw.core.exceptions import ExportError
from iil_researchfw.core.protocols import ResearchProjectProtocol

logger = logging.getLogger(__name__)


class ResearchExportService:
    """
    Export research results to various formats.

    Accepts any object implementing ResearchProjectProtocol.
    """

    def export_markdown(self, project: ResearchProjectProtocol) -> str:
        lines = [
            f"# {project.name}", "",
            f"**Query:** {project.query}",
            f"**Date:** {project.created_at.strftime('%Y-%m-%d')}",
            "",
        ]
        if project.description:
            lines += ["## Description", "", project.description, ""]
        findings = list(project.findings)
        if findings:
            lines += ["## Findings", ""]
            for f in findings:
                title = getattr(f, "title", getattr(f, "id", "Finding"))
                content = getattr(f, "content", "")
                source = getattr(f, "source_url", "")
                lines += [f"### {title}", content]
                if source:
                    lines.append(f"\n*Source: [{source}]({source})*")
                lines.append("")
        sources = list(project.sources)
        if sources:
            lines += ["## Sources", ""]
            for i, s in enumerate(sources, 1):
                url = getattr(s, "url", "")
                title = getattr(s, "title", url)
                lines.append(f"{i}. [{title}]({url})")
        return "\n".join(lines)

    def export_latex(self, project: ResearchProjectProtocol) -> str:
        date_str = project.created_at.strftime("%Y-%m-%d")
        lines = [
            r"\documentclass[12pt,a4paper]{article}",
            r"\usepackage[utf8]{inputenc}",
            r"\usepackage{hyperref}",
            r"\usepackage[backend=biber,style=apa]{biblatex}",
            "",
            rf"\title{{{self._latex_escape(project.name)}}}",
            rf"\date{{{date_str}}}",
            "", r"\begin{document}", r"\maketitle", "",
            r"\section*{Research Query}", self._latex_escape(project.query), "",
        ]
        findings = list(project.findings)
        if findings:
            lines += [r"\section*{Findings}", ""]
            for f in findings:
                title = getattr(f, "title", "Finding")
                content = getattr(f, "content", "")
                lines += [rf"\subsection*{{{self._latex_escape(title)}}}", self._latex_escape(content), ""]
        sources = list(project.sources)
        if sources:
            lines += [r"\section*{Sources}", r"\begin{enumerate}"]
            for s in sources:
                url = getattr(s, "url", "")
                title = getattr(s, "title", url)
                lines.append(rf"\item \href{{{url}}}{{{self._latex_escape(title)}}}")
            lines += [r"\end{enumerate}", ""]
        lines.append(r"\end{document}")
        return "\n".join(lines)

    def export_bibtex(self, project: ResearchProjectProtocol) -> str:
        entries = []
        for i, s in enumerate(project.sources):
            url = getattr(s, "url", "")
            title = getattr(s, "title", "")
            entries.append(
                f"@misc{{source{i + 1},\n"
                f"  title = {{{{{title}}}}},\n"
                f"  url = {{{url}}},\n"
                f"  year = {{{datetime.now().year}}},\n"
                f"}}"
            )
        return "\n\n".join(entries)

    def export_docx(self, project: ResearchProjectProtocol, output_path: str) -> str:
        """Export as DOCX. Requires iil-researchfw[export]."""
        try:
            from docx import Document  # type: ignore[import-untyped]
        except ImportError as exc:
            raise ExportError("python-docx not installed. Run: pip install iil-researchfw[export]") from exc

        doc = Document()
        doc.add_heading(project.name, 0)
        doc.add_paragraph(f"Query: {project.query}")
        doc.add_paragraph(f"Date: {project.created_at.strftime('%Y-%m-%d')}")
        if project.description:
            doc.add_heading("Description", level=1)
            doc.add_paragraph(project.description)
        findings = list(project.findings)
        if findings:
            doc.add_heading("Findings", level=1)
            for f in findings:
                doc.add_heading(getattr(f, "title", "Finding"), level=2)
                doc.add_paragraph(getattr(f, "content", ""))
        sources = list(project.sources)
        if sources:
            doc.add_heading("Sources", level=1)
            for i, s in enumerate(sources, 1):
                doc.add_paragraph(f"{i}. {getattr(s, 'title', '')} — {getattr(s, 'url', '')}")
        doc.save(output_path)
        return output_path

    def _latex_escape(self, text: str) -> str:
        for old, new in [
            ("\\", r"\textbackslash{}"), ("&", r"\&"), ("%", r"\%"),
            ("$", r"\$"), ("#", r"\#"), ("_", r"\_"), ("{", r"\{"),
            ("}", r"\}"), ("~", r"\textasciitilde{}"), ("^", r"\textasciicircum{}"),
        ]:
            text = text.replace(old, new)
        return text
